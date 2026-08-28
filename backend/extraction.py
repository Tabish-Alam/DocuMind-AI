"""
extraction.py
Pulls raw text out of uploaded documents (PDF, DOCX, TXT/MD).
"""
import io
from pypdf import PdfReader
from docx import Document as DocxDocument


class ExtractionError(Exception):
    pass


def extract_text(file_bytes: bytes, filename: str) -> str:
    """Route to the right extractor based on file extension."""
    lower = filename.lower()

    if lower.endswith(".pdf"):
        return _extract_pdf(file_bytes)
    elif lower.endswith(".docx"):
        return _extract_docx(file_bytes)
    elif lower.endswith((".txt", ".md")):
        return _extract_txt(file_bytes)
    else:
        raise ExtractionError(
            f"Unsupported file type for '{filename}'. "
            "Please upload a .pdf, .docx, .txt or .md file."
        )


def _extract_pdf(file_bytes: bytes) -> str:
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        pages = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            if text.strip():
                pages.append(f"[Page {i + 1}]\n{text.strip()}")
        full_text = "\n\n".join(pages)
        if not full_text.strip():
            raise ExtractionError(
                "No extractable text found in this PDF. It may be a scanned "
                "image without OCR — try a text-based PDF instead."
            )
        return full_text
    except ExtractionError:
        raise
    except Exception as exc:
        raise ExtractionError(f"Could not read PDF: {exc}") from exc


def _extract_docx(file_bytes: bytes) -> str:
    try:
        doc = DocxDocument(io.BytesIO(file_bytes))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]

        # also pull text out of tables
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

        full_text = "\n".join(parts)
        if not full_text.strip():
            raise ExtractionError("This DOCX file appears to be empty.")
        return full_text
    except ExtractionError:
        raise
    except Exception as exc:
        raise ExtractionError(f"Could not read DOCX: {exc}") from exc


def _extract_txt(file_bytes: bytes) -> str:
    for encoding in ("utf-8", "latin-1"):
        try:
            text = file_bytes.decode(encoding)
            if text.strip():
                return text
        except UnicodeDecodeError:
            continue
    raise ExtractionError("Could not decode text file.")
